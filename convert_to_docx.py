"""
Convert FINAL_REPORT.md to properly formatted DOCX for Gymnasiearbete submission.

This script uses python-docx to create a Word document with proper formatting
according to the Gymnasiearbete NA template requirements.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path


def setup_styles(doc):
    """Configure document styles to match Gymnasiearbete template."""
    
    # Normal text style
    styles = doc.styles
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(12)
    
    # Heading 1 style (e.g., "1. Inledning")
    try:
        style_h1 = styles['Heading 1']
    except KeyError:
        style_h1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    style_h1.font.name = 'Calibri'
    style_h1.font.size = Pt(14)
    style_h1.font.bold = True
    
    # Heading 2 style (e.g., "1.1 Syfte och mål")
    try:
        style_h2 = styles['Heading 2']
    except KeyError:
        style_h2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    style_h2.font.name = 'Calibri'
    style_h2.font.size = Pt(12)
    style_h2.font.bold = True
    
    # Heading 3 style
    try:
        style_h3 = styles['Heading 3']
    except KeyError:
        style_h3 = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    style_h3.font.name = 'Calibri'
    style_h3.font.size = Pt(12)
    style_h3.font.bold = True


def parse_markdown_to_docx(md_path, output_path, figures_dir):
    """
    Convert Markdown file to DOCX with proper formatting.
    
    Args:
        md_path: Path to FINAL_REPORT.md
        output_path: Path for output .docx file
        figures_dir: Path to figures directory
    """
    
    # Create document
    doc = Document()
    
    # Set up page margins (2.5 cm all around is standard)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Configure styles
    setup_styles(doc)
    
    # Read markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # State variables
    in_code_block = False
    in_table = False
    table_obj = None
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip horizontal rules
        if line.strip() in ['---', '___', '***']:
            i += 1
            continue
        
        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            # Add code as monospace
            p = doc.add_paragraph(line, style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            i += 1
            continue
        
        # Title (# Title)
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            # Main title - centered, large, bold
            if i < 10:  # First few lines are likely the main title
                p = doc.add_paragraph(title_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(18)
                    run.font.bold = True
            else:
                doc.add_heading(title_text, level=1)
            i += 1
            continue
        
        # Heading 2 (## Heading)
        if line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        # Heading 3 (### Heading)
        if line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue
        
        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line).strip()
            doc.add_paragraph(list_text, style='List Number')
            i += 1
            continue
        
        # Tables (simple markdown tables)
        if '|' in line and not in_table:
            # Start of table
            in_table = True
            # Parse header
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_obj = doc.add_table(rows=1, cols=len(cells))
            table_obj.style = 'Light Grid Accent 1'
            hdr_cells = table_obj.rows[0].cells
            for j, cell_text in enumerate(cells):
                hdr_cells[j].text = cell_text
                # Bold header
                for paragraph in hdr_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            i += 1
            # Skip separator line
            if i < len(lines) and '---' in lines[i]:
                i += 1
            continue
        
        if in_table and '|' in line:
            # Table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            row_cells = table_obj.add_row().cells
            for j, cell_text in enumerate(cells):
                row_cells[j].text = cell_text
            i += 1
            continue
        
        if in_table and '|' not in line:
            # End of table
            in_table = False
            table_obj = None
        
        # Images (figure references)
        if '*(Se `report_assets/figures/' in line:
            # Extract figure filename
            match = re.search(r'`report_assets/figures/(.*?)`', line)
            if match:
                fig_name = match.group(1)
                fig_path = figures_dir / fig_name
                if fig_path.exists():
                    # Add image
                    doc.add_picture(str(fig_path), width=Inches(6.0))
                    # Get the caption text (line before this one)
                    if i > 0:
                        caption_line = lines[i-1].strip()
                        if caption_line.startswith('**Figur'):
                            # Remove bold markers
                            caption = caption_line.replace('**', '')
                            p = doc.add_paragraph(caption, style='Caption')
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Bold text (**text**)
        if '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for idx, part in enumerate(parts):
                if idx % 2 == 0:
                    # Normal text
                    if part:
                        p.add_run(part)
                else:
                    # Bold text
                    run = p.add_run(part)
                    run.bold = True
            i += 1
            continue
        
        # Italic text (*text* or _text_)
        if '*' in line or '_' in line:
            # Handle inline emphasis (simplified)
            p = doc.add_paragraph(line)
            i += 1
            continue
        
        # Empty lines
        if not line.strip():
            # Only add paragraph break if previous wasn't empty
            if i > 0 and lines[i-1].strip():
                doc.add_paragraph()
            i += 1
            continue
        
        # Regular paragraph
        if line.strip():
            doc.add_paragraph(line.strip())
        
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"✓ Document saved to {output_path}")


def main():
    """Main conversion function."""
    
    # Check if python-docx is installed
    try:
        import docx
    except ImportError:
        print("ERROR: python-docx not installed.")
        print("Install with: pip install python-docx")
        return
    
    # Paths
    md_file = Path("report/FINAL_REPORT.md")
    output_file = Path("report/FINAL_REPORT.docx")
    figures_dir = Path("report_assets/figures")
    
    # Verify input file exists
    if not md_file.exists():
        print(f"ERROR: {md_file} not found")
        return
    
    print("Converting FINAL_REPORT.md to DOCX...")
    print(f"Input: {md_file}")
    print(f"Output: {output_file}")
    print(f"Figures: {figures_dir}")
    print()
    
    # Convert
    parse_markdown_to_docx(md_file, output_file, figures_dir)
    
    print()
    print("=" * 60)
    print("NEXT STEPS:")
    print("=" * 60)
    print("1. Open report/FINAL_REPORT.docx in Microsoft Word")
    print("2. Review formatting and adjust as needed:")
    print("   - Check page breaks before major sections")
    print("   - Verify figure placement and captions")
    print("   - Generate table of contents (References → Table of Contents)")
    print("   - Add page numbers (Insert → Page Number)")
    print("3. Add your supervisor's name on the title page")
    print("4. Save and submit!")
    print()


if __name__ == "__main__":
    main()
